from docx import Document
import os

doc = Document(r"C:\Users\leno2\Downloads\ola_cars_onboarding_workflow.docx")

output = []
for para in doc.paragraphs:
    if para.text.strip():
        # Include style info for headers
        style = para.style.name if para.style else ""
        if "Heading" in style:
            level = style.replace("Heading ", "").strip()
            try:
                level = int(level)
            except:
                level = 1
            output.append("#" * level + " " + para.text)
        else:
            output.append(para.text)

# Also extract tables
for table in doc.tables:
    output.append("\n--- TABLE ---")
    for row in table.rows:
        cells = [cell.text.strip() for cell in row.cells]
        output.append(" | ".join(cells))
    output.append("--- END TABLE ---\n")

result = "\n".join(output)

output_path = r"C:\Users\leno2\Desktop\OlaCarsBackend\tmp\docx_content.txt"
with open(output_path, "w", encoding="utf-8") as f:
    f.write(result)

print(f"Extracted {len(output)} lines to {output_path}")
